# app/api/routes/files.py
from __future__ import annotations

from typing import List  # NEU: für doc_ids-Liste im Bulk-Endpoint

from fastapi import (
    APIRouter,
    Depends,
    HTTPException,
    UploadFile,
    File,
    Form,
    Request,
    Query,
)
from starlette.responses import RedirectResponse, StreamingResponse
from sqlalchemy.orm import Session
from io import BytesIO
import os
import mimetypes

try:
    # Für DOCX-Text-Extraktion
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from app.core.config import settings  # falls anderswo genutzt
from app.utils.crypto_utils import decrypt_bytes, decrypt_text

from app.api.deps import get_current_user_web, CurrentUser
from app.db.database import get_db
from app.models.document import Document
from app.models.category import Category  # NEU: für Bulk-Zuweisung
from app.repositories.document_repo import get_document_owned
from app.services.version_service import (
    list_versions,
    upload_new_version,
    restore_version,
    rename_document_creates_version,
)
from app.services.text_similarity import compute_similarities
from app.web.templates import templates

router = APIRouter(prefix="/documents", tags=["documents"])


# -------------------------------------------------------------
# Helper: DOCX-Text extrahieren
# -------------------------------------------------------------
def extract_docx_text_from_bytes(data: bytes) -> str:
    """
    Extrahiert Text aus einer DOCX-Datei:
    - normale Absätze
    - Tabellenzellen (typisch für Rechnungs- / Vertragsvorlagen)

    Layout geht verloren, aber der reine Text wird lesbar dargestellt.
    """
    if DocxDocument is None:
        # python-docx nicht installiert
        print("[DOCX_DEBUG] DocxDocument is None (python-docx fehlt?)")
        return ""

    try:
        docx_file = DocxDocument(BytesIO(data))
    except Exception as e:
        print(f"[DOCX_DEBUG] Fehler beim Öffnen der DOCX-Datei: {e!r}")
        return ""

    parts: list[str] = []

    # 1) Normale Absätze
    for p in docx_file.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # 2) Tabellen (wichtig für Vorlagen / Formulare)
    for table in docx_file.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        parts.append(text)

    full_text = "\n\n".join(parts)
    print(
        f"[DOCX_DEBUG] extrahierte Blöcke: {len(parts)}, "
        f"Zeichen insgesamt: {len(full_text)}"
    )
    return full_text


# -------------------------------------------------------------
# DETAIL-SEITE (Metadaten + Versionen)
# -------------------------------------------------------------
@router.get("/{doc_id}")
def detail(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Zeigt Metadaten und Versionsliste eines Dokuments.
    Nutzt das Template document_detail.html.
    """
    doc = get_document_owned(db, doc_id, user.id)
    if not doc:
        raise HTTPException(status_code=404, detail="Not found")

    _, versions = list_versions(db, user.id, doc_id)

    return templates.TemplateResponse(
        "document_detail.html",
        {
            "request": request,
            "user": user,
            "document": doc,
            "versions": versions,
        },
    )


# -------------------------------------------------------------
# VIEWER-SEITE (PDF im pdf.js-Viewer oder Text/OCR/DOCX)
# -------------------------------------------------------------
@router.get("/{doc_id}/view")
def document_view_page(
    doc_id: int,
    request: Request,
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Viewer-Seite für ein Dokument.

    - PDFs:
        werden über /documents/{id}/download?inline=1 im Browser angezeigt
        (die HTML-Seite nutzt pdf.js, um alle Seiten zu rendern)
    - DOCX/TXT:
        Text wird aus Datei extrahiert (falls kein OCR-Text vorhanden)
    - andere:
        ggf. OCR-Text anzeigen
    """
    # Nur Dokumente des eingeloggten Users
    doc = get_document_owned(db, doc_id, user.id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Datei laden (für Text-/DOCX-Vorschau und ggf. PDF-Header-Check)
    file_bytes: bytes | None = None
    storage_path = getattr(doc, "storage_path", None)
    if storage_path and os.path.exists(storage_path):
        try:
            with open(storage_path, "rb") as f:
                encrypted_bytes = f.read()
            file_bytes = decrypt_bytes(encrypted_bytes)
        except Exception:
            file_bytes = None

    # Dateiname für Anzeige / Typ-Erkennung
    raw_name = (
        getattr(doc, "original_filename", None)
        or getattr(doc, "filename", None)
        or getattr(doc, "name", None)
        or ""
    ).strip()
    name_lower = raw_name.lower()

    # MIME-Typ robust bestimmen (DB → Dateiname → Fallback)
    mime_from_db = getattr(doc, "mime_type", None)
    mime_from_name, _ = mimetypes.guess_type(raw_name)
    mime_raw = mime_from_db or mime_from_name or "application/octet-stream"
    # charset etc. abschneiden, alles kleinschreiben
    mime_norm = mime_raw.split(";")[0].strip().lower()

    # ---------------------------------------------------------
    # PDF-Erkennung:
    # 1) MIME "application/pdf"
    # 2) Dateiendung .pdf
    # 3) Magic-Header %PDF-
    # ---------------------------------------------------------
    is_pdf = False

    if mime_norm == "application/pdf" or name_lower.endswith(".pdf"):
        is_pdf = True

    if (not is_pdf) and file_bytes:
        head = file_bytes.lstrip()[:5]
        if head.startswith(b"%PDF-") or head.startswith(b"%pdf-"):
            is_pdf = True

    # DOCX-Erkennung (nur wenn es kein PDF ist)
    is_docx = (not is_pdf) and name_lower.endswith(".docx")

    # ---------------------------------------------------------
    # Klartext / OCR-Text
    # ---------------------------------------------------------
    ocr_text: str | None = None

    # 1) OCR-Text aus der DB entschlüsseln (falls vorhanden)
    if getattr(doc, "ocr_text", None):
        try:
            ocr_text = decrypt_text(doc.ocr_text)
        except Exception:
            ocr_text = "[Fehler beim Entschlüsseln des OCR-Texts]"

    # 2) Falls kein OCR-Text, aber Datei vorhanden und kein PDF → Text aus Datei
    if (not is_pdf) and (not ocr_text) and file_bytes:
        try:
            if name_lower.endswith(".docx"):
                # Word-Dokument: Fließtext + Tabellen
                text = extract_docx_text_from_bytes(file_bytes)
                if text:
                    ocr_text = text
            elif name_lower.endswith(".txt"):
                # Einfache Textdatei
                text = file_bytes.decode("utf-8", errors="replace")
                if text.strip():
                    ocr_text = text
        except Exception:
            # Bei Fehlern einfach keinen Text anzeigen
            pass

    # Debug-Ausgabe in der Konsole (hilfreich zum Prüfen)
    print(
        f"[VIEW_DEBUG] id={doc.id}, name={raw_name!r}, mime_raw={mime_raw!r}, "
        f"mime_norm={mime_norm!r}, is_pdf={is_pdf}, is_docx={is_docx}, "
        f"has_file_bytes={file_bytes is not None}, has_ocr={ocr_text is not None}"
    )

    # Template bekommt Informationen, ob PDF oder DOCX,
    # plus optionalen Klartext (OCR oder extrahierter Text)
    return templates.TemplateResponse(
        "document_view.html",
        {
            "request": request,
            "user": user,
            "doc": doc,
            "is_pdf": is_pdf,
            "is_docx": is_docx,
            "ocr_text": ocr_text,
        },
    )


# -------------------------------------------------------------
# DOWNLOAD + INLINE-VORSCHAU
# -------------------------------------------------------------
@router.get("/{doc_id}/download")
def download(
    doc_id: int,
    inline: bool = Query(False),
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Liefert die Datei:
    - inline=true  → Content-Disposition: inline  (für iframe / pdf.js-Vorschau)
    - inline=false → Content-Disposition: attachment (klassischer Download)
    """
    doc = get_document_owned(db, doc_id, user.id)
    if not doc:
        raise HTTPException(status_code=404, detail="Not found")

    if not doc.storage_path or not os.path.exists(doc.storage_path):
        raise HTTPException(status_code=404, detail="File not found on server")

    with open(doc.storage_path, "rb") as f:
        encrypted_bytes = f.read()

    try:
        decrypted = decrypt_bytes(encrypted_bytes)
    except Exception:
        raise HTTPException(status_code=500, detail="Failed to decrypt file")

    stream = BytesIO(decrypted)

    download_name = (
        getattr(doc, "original_filename", None)
        or getattr(doc, "filename", None)
        or f"document-{doc.id}"
    )

    mime = (
        getattr(doc, "mime_type", None)
        or mimetypes.guess_type(download_name)[0]
        or "application/octet-stream"
    )

    disposition = "inline" if inline else "attachment"

    return StreamingResponse(
        stream,
        media_type=mime,
        headers={
            "Content-Disposition": f'{disposition}; filename="{download_name}"',
            "X-Content-Type-Options": "nosniff",
        },
    )


# -------------------------------------------------------------
# VERSIONEN / RENAME
# -------------------------------------------------------------
@router.post("/{doc_id}/rename-web")
def rename(
    doc_id: int,
    new_title: str = Form(...),
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Benennt ein Dokument um und legt dabei eine neue Version an.
    """
    try:
        rename_document_creates_version(db, user.id, doc_id, new_title.strip())
        return RedirectResponse(url=f"/documents/{doc_id}", status_code=303)
    except ValueError:
        raise HTTPException(status_code=404, detail="Not found")


@router.post("/{doc_id}/upload-version-web")
def upload_version(
    doc_id: int,
    file: UploadFile = File(...),
    note: str = Form(""),
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Lädt eine neue Version zu einem bestehenden Dokument hoch.
    """
    try:
        upload_new_version(db, user.id, doc_id, file.filename, file.file, note=note)
        return RedirectResponse(url=f"/documents/{doc_id}", status_code=303)
    except ValueError:
        raise HTTPException(status_code=404, detail="Not found")


@router.post("/{doc_id}/restore-web")
def restore(
    doc_id: int,
    version_id: int = Form(...),
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Stellt eine ältere Version eines Dokuments wieder her.
    """
    try:
        restore_version(db, user.id, doc_id, version_id)
        return RedirectResponse(url=f"/documents/{doc_id}", status_code=303)
    except ValueError:
        raise HTTPException(status_code=404, detail="Not found")


# -------------------------------------------------------------
# BULK: Mehrere Dokumente einer Kategorie zuweisen
# -------------------------------------------------------------
@router.post("/bulk-assign-category")
def bulk_assign_category(
    category_id: int = Form(...),
    doc_ids: List[int] = Form(...),
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Weist mehrere Dokumente (doc_ids) auf einmal einer Kategorie zu.

    Annahmen:
    - category_id gehört zum eingeloggten User
    - nur Dokumente des eingeloggten Users werden geändert
    """

    # Kategorie validieren (nur Kategorien des Users)
    category = (
        db.query(Category)
        .filter(
            Category.id == category_id,
            Category.user_id == user.id,
        )
        .first()
    )
    if not category:
        raise HTTPException(status_code=404, detail="Category not found")

    # Sicherstellen, dass überhaupt IDs angekommen sind
    if not doc_ids:
        # Kein Fatal Error → zurück zur Liste
        return RedirectResponse(url="/documents?bulk=none", status_code=303)

    # Dokumente des Users mit den IDs holen
    docs = (
        db.query(Document)
        .filter(
            Document.id.in_(doc_ids),
            Document.owner_user_id == user.id,
            Document.is_deleted == False,
        )
        .all()
    )

    if not docs:
        return RedirectResponse(url="/documents?bulk=nodocs", status_code=303)

    # Kategorie setzen
    for d in docs:
        d.category_id = category_id

    db.commit()

    # Zurück zur Übersicht mit kleinem Query-Flag für späteres Feedback
    return RedirectResponse(url="/documents?bulk=ok", status_code=303)


# -------------------------------------------------------------
# ÄHNLICHE DOKUMENTE (OCR)
# -------------------------------------------------------------
@router.get("/{doc_id}/similar")
def similar_documents(
    doc_id: int,
    db: Session = Depends(get_db),
    user: CurrentUser = Depends(get_current_user_web),
):
    """
    Einfache Ähnlichkeitssuche auf Basis des entschlüsselten OCR-Textes.
    """
    docs = (
        db.query(Document)
        .filter(
            Document.owner_user_id == user.id,
            Document.is_deleted == False,
            Document.ocr_text.isnot(None),
        )
        .order_by(Document.id.asc())
        .all()
    )

    id_to_index = {d.id: i for i, d in enumerate(docs)}
    if doc_id not in id_to_index:
        raise HTTPException(
            status_code=404,
            detail="Document not found or no OCR text",
        )

    corpus: list[str] = []
    for d in docs:
        if d.ocr_text:
            corpus.append(decrypt_text(d.ocr_text))
        else:
            corpus.append("")

    idx = id_to_index[doc_id]
    sims = compute_similarities(corpus, idx, top_k=5)

    result = [
        {
            "id": docs[i].id,
            "filename": docs[i].filename,
            "score": score,
        }
        for i, score in sims
    ]

    return {"base_id": doc_id, "similar": result}
