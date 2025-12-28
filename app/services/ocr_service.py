# app/services/ocr_service.py
from __future__ import annotations

import os

from app.services.ocr_analysis import clean_text, run_ocr_on_file


def extract_text_from_pdf(path: str, lang: str = "deu+eng", dpi: int = 300) -> str:
    # Erst versuchen: Text-Layer (falls vorhanden), sonst OCR-Fallback
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        PdfReader = None  # type: ignore

    if PdfReader is not None:
        try:
            reader = PdfReader(path)
            text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
            if text:
                return text
        except Exception:
            pass

    return run_ocr_on_file(path, lang=lang, dpi=dpi)


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document as DocxDocument  # type: ignore
    except Exception as exc:
        raise RuntimeError(f"python-docx fehlt oder kann nicht importiert werden: {exc!r}") from exc

    doc = DocxDocument(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)

    return "\n".join(parts).strip()


# NEU: fehlt, wird von debug_ocr.py importiert
def extract_text_from_image(path: str, lang: str = "deu+eng") -> str:
    # raw OCR (ohne clean), damit debug route exakt zeigt, was OCR liefert
    return run_ocr_on_file(path, lang=lang)


def ocr_and_clean(path: str, lang: str = "deu+eng", dpi: int = 300) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        raw = extract_text_from_pdf(path, lang=lang, dpi=dpi)
        return clean_text(raw)

    if ext == ".docx":
        raw = extract_text_from_docx(path)
        return clean_text(raw)

    if ext in {".txt", ".md", ".csv", ".log"}:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return clean_text(f.read())

    raw = run_ocr_on_file(path, lang=lang, dpi=dpi)
    return clean_text(raw)
